from turtle import color
import requests  #用於get請求
from bs4 import BeautifulSoup as bs #網頁分析
from os import mkdir #建資料夾(目錄)
import os #用於檔案/資料夾的操作
import subprocess #打開資料夾

#GUI
import tkinter as tk
from tkinter import messagebox

import re
from docx import Document
from docx.shared import Pt, Cm #處理大小，字體或空間
from docx.oxml.ns import qn
from PIL import Image           #處理圖片規格

import unicodedata #全形轉半形

#從網頁抓取小說
def getReq(url): #傳入 
    # url = 'https://www.2dfan.com/topics/10046'
    # url = 'https://www.2dfan.com/topics/10046/page/8'
    
    #可能需改動項目1
    cookies = {
        "__bid_n": "186c15da74c74505484207",
        "remember_me_token": "eyJfcmFpbHMiOnsibWVzc2FnZSI6IkluQXRZMEo2YzBGQmRWaDFWamQwVjJzMVVWWllJZz09IiwiZXhwIjoiMjAyMy0wNy0zMFQxNjozODo1NS44NzdaIiwicHVyIjpudWxsfX0%3D--5d28c2540a857ba1cd45e816048812a089b4a65c",
        "Hm_lvt_79251201618e9337c1169fef9b3e4786": "1682702013,1682777534,1682903386,1683982331",
        "_project_hgc_session": "dWxlYmNLODUySHVZUWczQW1nb3FYeGZzb2MzZGxqc0JKS0RiWWVWRWkrNHcyQmRiOTZzRmUwTlBWcTRBcURoYU4yeG4xWE44L0Jmc2ZhL0QvZjA4ZnE1NXh0RWhtd1FtYnRoUktsUUhKSmR4ZnRmM1YwT2UrSVNXSDVlcnd1emhZUktrL3NwWVovMjlyQlJEcGJGNFFPemgwMk1wRkE5K1Vib2lTUWZKZHB1ZitkVWVld1VwU0dWSjBucmp1OVlFLS1hR2t3WmhZVG93RHF1YXpOWmdERDZRPT0%3D--0a5aa43cc538a4c70b95b11c48f7e8a7303e1e91",
        "Hm_lpvt_79251201618e9337c1169fef9b3e4786": "1684244906"
    }

    headers={ #默認
        'Accept':'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7',
        'Upgrade-Insecure-Requests':'1',#不用改
        'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/112.0.0.0 Safari/537.36 Edg/112.0.1722.39',
        # 'Connection': 'keep-alive' #保持連接用
        "Referer":'https://www.2dfan.com/',
    }  
    req = requests.get(url, headers=headers, cookies=cookies, verify=False)
    if req.status_code != 200:
        print("連接失敗")
    return req

def getSoup(req):
    soup = bs(req.text, "html.parser")
    return soup

##取得頁數##
def getPage(soup):
    pageInt = 1 #預設只有一頁
    #若有多頁，從'尾页'連結上取得總頁數
    page = soup.find('div', {'id': 'content-pagination'})
    pageText = page.text.replace('\n', '')
    # pageText = removeSpaces(pageText)
    if (page != None) and (pageText != ''):     
        print("######")  
        nextPage = page.find('a', text='尾页')
        nextPageLinks= nextPage['href']

        # 篩出連結末端的頁數
        page_pattern = r'[0-9]+$'
        pageInt = int(re.search(page_pattern, nextPageLinks).group())

    return pageInt


def getContentList(url, pageInt):
    result = []
    #逐次依頁數取出
    for i in range(pageInt):
        #不是第一頁的話要加上後綴
        if pageInt != 1:
            soup = getSoup( getReq( url + '/page/' + str(i+1) ) )
        else:
            soup = getSoup( getReq(url) )
        #移除指定標籤(soup, 標籤名, 指定屬性, 指定屬性內容)
        removeTag(soup, 'ul', 'class', 'breadcrumb')

        
        # 找到指定的 div 並取其所有標籤
        div = soup.find('div', {'id': 'topic-content'})
        tags = div.find_all() # get html tags
        
        result.append (tags)

        # print("############## 第"+ str(i+1) +"頁取出 ##############")
        # print(tags)
    return result
 
#消除前後空白
def removeSpaces(str):
    return  re.sub(r'(^\s*)|(\s*$)', "", str)


#建立word並全域性設定【邊界(窄)】及【字型】
def newDocx():
    document = Document()
    section = document.sections[0] # 獲取頁面節點
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    #A4紙張大小
    section.page_height = Cm(29.7)  # 高
    section.page_width = Cm(21.0)   # 寬

    document.styles['Normal'].font.name = u'微軟正黑體'  # 字型
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')  # 中文字型需再新增這個設定
    document.styles['Normal'].font.size = Pt(12)  # 字號 四號對應14
    return document

#寫入 word
def editDocx(url, req, soup, document, tagsList):
    
    #標註來源
    document.add_paragraph( '來源：' +url)
    #逐頁處理 list
    n = 0
    for tags in tagsList:
        n += 1
        print("### 第 "+str(n)+" 頁 ###")
        #逐行處理html標籤
        for tag in tags:
            #處理文本
            # print(tag.text)
            if tag.name != 'img':
                #跳過
                if tag.text[:4] == "开场动画" or tag.text == '':   
                    continue
                #不需要末端的預覽圖
                if tag.text == "画廊欣赏" or tag.text[:2] == 'CG': 
                    break
        
                context = strReplace(tag.text)         #字串處理
                if (tag.name == 'p'):   
                    context += '\n'   #<p>尾端加上換行
                paragraph = document.add_paragraph(context)

                #段落改動，若是標題則套用標題樣式
                if  tag.name == 'h4':
                    paragraph.style = document.styles['Heading 2']
                    paragraph.style.font.size = Pt(18) #字體大小
                else:
                    paragraph.style.font.size = Pt(12) #字體大小
                    #段落勾選"相同樣式的個段落之間不要加上空間"
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(0)
                    paragraph_format.space_after = Pt(0)
                    #改行距
                    paragraph_format.line_spacing = Pt(18)
                #開頭跟結尾是對話框的話，改為粗體
                if context[:1] == '「' and context[len(context)-2:len(context)] == '」\n':
                    paragraph.runs[0].bold = True 
                #<p>末尾加上空行
                    
            #處理圖片
            else: 
                #取圖片連結
                # print("#### 圖片 ####")
                image_url = tag.get('src')
                
                # 读取图片并保存到本地
                #保存图片至本地
                img_name = 'tmpPic.jpg'
                with open(img_name,'wb')as f:
                    req = getReq(image_url)
                    response = req.content
                    f.write(response)
                    f.close()

                # 插入圖片，並調整大小 (document, 圖片路徑)
                insert_image_with_resizing(document, 'tmpPic.jpg')
                
    #介紹標題
    titleName = soup.select('.block > div > h3')[0].text
    #完整檔名
    fileName = toFileName(titleName + ".docx") 
    #取得docx完整路徑
    docx_path = os.path.abspath(fileName)
    # 保存 Word 文档
    document.save(fileName)
    os.remove("tmpPic.jpg")
    print("Word 已保存")
    # 使用subprocess打開資料夾，並顯示在最上層
    subprocess.Popen('explorer /select,"' + docx_path + '"')
    # https://galge.fun/subjects/9793   主頁
    # https://galge.fun/topics/12649    介紹頁

#GUI
def new_window():  
    window = tk.Tk()
    window.title('新視窗')
    window.geometry('400x150') #寬x高
    window.maxsize(1280,700)

    #元件類別(父類別, 選擇性參數1 = 值1, ...) ，建立元件
    #元件.grid(row=列數, column=行數) ，設定(相對)位置
    mylabel = tk.Label(window, text='請輸入網址：', height = 2, width=12)
    mylabel.grid(row=0, column=0)

    intput_website =tk.StringVar()
    myEntry = tk.Entry(window, textvariable=intput_website, width=40) #文字輸入框
    myEntry.grid(row=0, column=1)
    myEntry.focus() #放入游標

    mylabel_start = tk.Label(window, text='')
    mylabel_start.grid(row=2, column=0, columnspan=10, sticky="w") #橫跨10欄 靠W(西)對齊

    # myText = tk.Text(window, width=40, height=10) #文字框 height=行數
    # myText.grid(row=3, column=0, columnspan=50)

    #button觸發event
    def button_event():
        if intput_website.get() == '':
            tk.messagebox.showerror('message', '未輸入') #showerror 提醒訊息
        global  output_website 
        webURL = intput_website.get()

        mylabel_start.config(text='開始抓取...')
        # myText.delete("1.0","end") #先清空

        ###############################
        req = getReq(webURL)       #requests物件
        soup = getSoup(req)     
        page = getPage(soup)    #取得頁數
        tagsList = getContentList( webURL , page )
        #創建word檔
        document = newDocx()
        #開始放入內容
        editDocx( webURL, req, soup, document, tagsList ) 
        ###############################

        mylabel_start.config(text='抓取完畢!!')


    myButton = tk.Button(window, text='執行', command= button_event, height = 5, width=12, background='orange')
    myButton.grid(row=1, column=0, columnspan=4)
    #inp1 = tk.Entry(window, text="Hello World", bg="yellow", fg="#263238", font=('Arial', 20))，進一步元件風格
    
    window.mainloop()

#替換
def strReplace(text):
    text = text.replace('\n','')    #拿掉換行
    text = text.replace('“', '「').replace('”', '」')   #對話框替換
    return text

#檢查特殊字元，用全形取代特殊字元，避免不符檔名規則
def toFileName(string):
    string = full2half(string)            #全形轉半形
    string = re.sub(r'^\s+?',"",string)   #消除前面空白
    flag = False
    for i in string: #逐一比對
        if i=='\\' or i=='/' or i==':' or i=='*' or i=='?' or i=='<' or i=='>' or i=='|':
            flag = True
    if flag: #替換成全形
        string = string.replace('\\','＼').replace('/','／').replace(':','：')
        string = string.replace('*','＊').replace('?','？').replace('<','＜')
        string = string.replace('>','＞').replace('|','｜')
    return string.replace('―','—')
    # return string

#全形轉半形
def full2half(c: str) -> str:
    return unicodedata.normalize("NFKC", c)

#從Soup移除指定Tag
#(soup物件, 標籤名, 屬性名, 屬性內容)
def removeTag(soup, tagName, key, value):
        # 找到所有指定標籤(tag)
        tags = soup.find_all(tagName)

        # 迭代每個標籤(tag)
        for tag in tags:
            if key != '':
                # 如果標籤(tag)有屬性(key)為指定內容(value)
                if value in tag.get(key, []):
                    # 移除該標籤
                    tag.extract()
            #若不指定屬性 (key == '')，直接移除標籤(tag)
            else: 
                tag.extract()

def insert_image_with_resizing(doc, image_path):
    
    # A4編輯區長寬
    max_height = Cm(27.16)
    max_width = Cm(18.46)

    # 取圖片的長寬
    img = Image.open(image_path)
    width, height = img.size
    # 像素轉cm
    width = Cm(width / 28.35)
    height = Cm(height/ 28.35)
    # 計算長寬比例
    scale = width/height
    
    # 若圖片超過邊界，縮小至符合邊界
    # 不考慮的情境：寬度符合邊界後，高度仍超過邊界
    if max_width < width:       #寬度符合邊界
        new_height = max_width / scale
        new_width = max_width
    elif max_height < height:   #高度符合邊界
        new_height = max_height
        new_width = max_height * scale
    else:                       #不變
        new_height = height
        new_width = width
    
    # 插入圖片並設置長寬
    doc.add_picture(image_path, width=new_width, height=new_height)


if __name__=='__main__':
    new_window()
