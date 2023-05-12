from mimetypes import MimeTypes
from turtle import width
from docx import Document
import os
import tkinter as tk
from docx.shared import Pt #插入圖片  Floating point length
import re



#讀取內容
# print("段落數量：", len(document.paragraphs))
# for para in document.paragraphs:
#     print(para.text)


#傳入目標路徑，取得目錄下所有檔名(不包括子資料夾及其檔案)的完整路徑
def getAllFileNameList(dirPath):
    fileList = [dirPath+f for f in os.listdir(dirPath) if os.path.isfile(os.path.join(dirPath, f))]
    return fileList

#傳入目標路徑，取得目錄下所有檔名(不包括子資料夾及其檔案)
def getFileNameList(dirPath):
    fileList = [f for f in os.listdir(dirPath) if os.path.isfile(os.path.join(dirPath, f))]
    return [fileList]


#傳入txt檔名list
def writeInWord(txtDirPathList,picDirPath):
    pattern = re.compile(r'([^\s,]+?\.jpg)',re.I)#建立圖名匹配pattern
    pattern2 = re.compile(r'bg',re.I)#建立圖名匹配pattern
    pattern3 = re.compile(r'view',re.I)#建立圖名匹配pattern

    for txtfileName in txtDirPathList: #逐個開啟txt檔
        print('檔名為：' + txtfileName)
        #Word檔案操作
        document = Document() #新開並建立docx檔
        with open(txtfileName,'r',encoding='utf-16-le') as f:
            for line in f:  #逐行
                #字串取代，取代掉無用的圖名，避免一行有兩圖名
                line = line.replace('<KW><WinClear>','').replace('\n','')
                line = line.replace("black.bmp",'').replace('white.bmp','')
                line = line.replace('から.bmp','')
                line = line.replace('.bmp','.jpg')

                #找圖片位置(正則表達式，對每行進行比對並插入)
                find = pattern.findall(line) #返回匹配的圖名(物件)
                
                if find != []:
                    picPath = picDirPath + find[0]#圖片相對路徑
                    # print(picPath)
                    isfile = os.path.isfile(picPath) #檢查檔案是否存在
                    if isfile:
                        if pattern2.match(find[0]) or pattern3.match(find[0]): #判斷是否背景圖
                            document.add_picture(picPath,width=Pt(300))# 插入圖片
                        else:
                            document.add_picture(picPath)# 插入圖片
                # 寫入docx
                document.add_paragraph(line)#以段落方式寫入到docx
                
        #以原txt檔名保存
        txtfileNameDoc = txtfileName.replace(txtDirPath,'').replace('.txt','') 
        document.save('AtelierGH/DOCX/' + txtfileNameDoc +'.docx') #保存

#GUI 暫時沒用到
def new_window():  
    window = tk.Tk()
    window.title('新視窗')
    window.geometry('300x100') #寬x高
    window.maxsize(1280,700)

    #元件類別(父類別, 選擇性參數1 = 值1, ...) ，建立元件
    #元件.grid(row=列數, column=行數) ，設定(相對)位置
    mylabel = tk.Label(window, text='請輸入網址：')
    mylabel.grid(row=0, column=0)

    intput_website =tk.StringVar()
    myEntry = tk.Entry(window, textvariable=intput_website, width=30) #文字輸入框
    myEntry.grid(row=0, column=1)
    myEntry.focus() #放入游標

    mylabel_start = tk.Label(window, text='')
    mylabel_start.grid(row=2, column=0)

    #button觸發event
    def button_event():
        if intput_website.get() == '':
            tk.messagebox.showerror('message', '未輸入') #showerror 提醒訊息
        else:
            #取得文字框內容
            global  output_website 
            output_website = intput_website.get()

            mylabel_start.config(text='正在運作...')
            mylabel_start.config(text='下載完畢!!')

    myButton = tk.Button(window, text='計算', command= button_event)
    myButton.grid(row=1, column=0)
    #inp1 = tk.Entry(window, text="Hello World", bg="yellow", fg="#263238", font=('Arial', 20))，進一步元件風格
    
    window.mainloop()


if __name__ == '__main__':
    #文字資料夾
    txtDirPath = 'AtelierGH/TXT/'
    txtDirPathList = getAllFileNameList(txtDirPath)
    #圖片資料夾
    picDirPath = 'AtelierGH/JPG/allJPG/'
    picNameList = getFileNameList(picDirPath)

    writeInWord(txtDirPathList,picDirPath) #參數 文檔list,圖檔list,圖檔路徑
    # new_window()